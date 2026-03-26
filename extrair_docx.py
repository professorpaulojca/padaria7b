"""
Script para extrair texto dos documentos .docx originais
e salvar como Markdown em docs/requisitos/00-originais/ e 01-normalizados/
"""

import os
from pathlib import Path
from docx import Document

WORKSPACE = Path(__file__).parent
DOCS_DIR = WORKSPACE / "docs"
ORIGINAIS_DIR = DOCS_DIR / "requisitos" / "00-originais"
NORMALIZADOS_DIR = DOCS_DIR / "requisitos" / "01-normalizados"

# Criar pastas se não existirem
ORIGINAIS_DIR.mkdir(parents=True, exist_ok=True)
NORMALIZADOS_DIR.mkdir(parents=True, exist_ok=True)


def extrair_docx(caminho_docx: Path) -> str:
    """Extrai texto de um arquivo .docx preservando estrutura de parágrafos."""
    doc = Document(str(caminho_docx))
    linhas = []

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            linhas.append("")
            continue

        style = para.style.name.lower() if para.style else ""

        if "heading 1" in style or "título 1" in style:
            linhas.append(f"# {texto}")
        elif "heading 2" in style or "título 2" in style:
            linhas.append(f"## {texto}")
        elif "heading 3" in style or "título 3" in style:
            linhas.append(f"### {texto}")
        elif "heading" in style or "título" in style:
            linhas.append(f"#### {texto}")
        elif "list" in style or "bullet" in style:
            linhas.append(f"- {texto}")
        else:
            linhas.append(texto)

    # Extrair tabelas
    for i, tabela in enumerate(doc.tables):
        linhas.append("")
        linhas.append(f"### Tabela {i + 1}")
        for j, row in enumerate(tabela.rows):
            celulas = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            linha_md = "| " + " | ".join(celulas) + " |"
            linhas.append(linha_md)
            if j == 0:
                separador = "| " + " | ".join(["---"] * len(celulas)) + " |"
                linhas.append(separador)

    return "\n".join(linhas)


def main():
    # Buscar todos os .docx em docs/
    docx_files = list(DOCS_DIR.glob("*.docx"))

    if not docx_files:
        print("Nenhum arquivo .docx encontrado em docs/")
        return

    for docx_path in docx_files:
        nome_base = docx_path.stem
        print(f"\n{'='*60}")
        print(f"Processando: {docx_path.name}")
        print(f"{'='*60}")

        # Extrair conteúdo
        conteudo_md = extrair_docx(docx_path)

        # Salvar em 00-originais (texto bruto extraído)
        arquivo_original = ORIGINAIS_DIR / f"{nome_base}.md"
        with open(arquivo_original, "w", encoding="utf-8") as f:
            f.write(f"# {nome_base}\n\n")
            f.write(f"> Fonte: `docs/{docx_path.name}`\n\n")
            f.write(conteudo_md)
        print(f"  -> Salvo em: {arquivo_original.relative_to(WORKSPACE)}")

        # Salvar em 01-normalizados (mesmo conteúdo por enquanto, será revisado)
        arquivo_normalizado = NORMALIZADOS_DIR / f"{nome_base}.md"
        with open(arquivo_normalizado, "w", encoding="utf-8") as f:
            f.write(f"# {nome_base}\n\n")
            f.write(f"> Fonte: `docs/{docx_path.name}`  \n")
            f.write(f"> Status: **Extração automática — requer revisão manual**\n\n")
            f.write(conteudo_md)
        print(f"  -> Salvo em: {arquivo_normalizado.relative_to(WORKSPACE)}")

        # Imprimir conteúdo no terminal para verificação
        print(f"\n--- CONTEÚDO EXTRAÍDO ({docx_path.name}) ---")
        print(conteudo_md)
        print(f"--- FIM ({docx_path.name}) ---\n")

    print(f"\nTotal de documentos processados: {len(docx_files)}")
    print(f"Arquivos em 00-originais: {ORIGINAIS_DIR}")
    print(f"Arquivos em 01-normalizados: {NORMALIZADOS_DIR}")


if __name__ == "__main__":
    main()
